# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# 辅助方法：设置单元格背景色
# ==========================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

# ==========================================
# 辅助方法：设置单元格内边距
# ==========================================
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# ==========================================
# 辅助方法：设置表格边框（细灰边框）
# ==========================================
def set_table_borders(table, color="D3D3D3"):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # ~0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ==========================================
# 辅助方法：合并单元格并设置文本
# ==========================================
def merge_cells_and_set_text(table, start_row, start_col, end_row, end_col, text, align_center=False, bold=False, size=10.5, font_name="SimSun"):
    # 合并单元格
    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)
    merged_cell = start_cell.merge(end_cell)
    
    # 清理并设置文本
    merged_cell.text = ""
    p = merged_cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.ascii_name = 'Times New Roman'
    run.font.hansi_name = font_name
    
    return merged_cell

# ==========================================
# 辅助方法：向段落中添加带有粗体和行内代码的文本
# ==========================================
def add_paragraph_with_runs(p, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_text = part[2:-2]
            run = p.add_run(clean_text)
            run.bold = True
            run.font.name = 'SimSun'
            run.font.ascii_name = 'Times New Roman'
            run.font.hansi_name = 'SimSun'
        elif part.startswith('`') and part.endswith('`'):
            clean_sub = part[1:-1]
            run = p.add_run(clean_sub)
            run.font.name = 'Consolas'
            run.font.ascii_name = 'Consolas'
            run.font.hansi_name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78) # 行内代码红色高亮
        else:
            run = p.add_run(part)
            run.font.name = 'SimSun'
            run.font.ascii_name = 'Times New Roman'
            run.font.hansi_name = 'SimSun'

# ==========================================
# 核心生成类
# ==========================================
class ReportGenerator:
    def __init__(self, md_path, docx_path):
        self.md_path = md_path
        self.docx_path = docx_path
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        # 页面尺寸：A4
        for section in self.doc.sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            # 常规页边距：上2.54cm，下2.54cm，左3.18cm，右3.18cm
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def generate_cover_page(self):
        """生成封面页"""
        # 添加校名标题
        p_school = self.doc.add_paragraph()
        p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_school.paragraph_format.space_before = Pt(40)
        p_school.paragraph_format.space_after = Pt(20)
        run_school = p_school.add_run("河 南 工 业 大 学")
        run_school.font.name = 'SimHei'
        run_school.font.hansi_name = 'SimHei'
        run_school.font.size = Pt(28) # 二号
        run_school.bold = True
        run_school.font.color.rgb = RGBColor(192, 0, 0) # 喜庆红校名

        # 间隔
        p_space = self.doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(60)

        # 课程设计主标题
        p_title1 = self.doc.add_paragraph()
        p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title1 = p_title1.add_run("《软件开发综合实践》")
        run_title1.font.name = 'SimHei'
        run_title1.font.hansi_name = 'SimHei'
        run_title1.font.size = Pt(24) # 小一
        run_title1.bold = True

        p_title2 = self.doc.add_paragraph()
        p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title2.paragraph_format.space_after = Pt(80)
        run_title2 = p_title2.add_run("课程设计报告")
        run_title2.font.name = 'SimHei'
        run_title2.font.hansi_name = 'SimHei'
        run_title2.font.size = Pt(24)
        run_title2.bold = True

        # 信息对齐表格 (3.5 英寸宽，居中，无边框)
        table = self.doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # 设置列宽
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(3.7)

        properties = [
            ("题      目：", "基于Go语言与云原生微服务架构的社交平台的设计与实现"),
            ("专业班级：", "软件 230X 班"),
            ("学生姓名：", ""),
            ("学      号：", ""),
            ("指导教师：", ""),
            ("课程设计时间：", "2026.6.3—2026.7.1")
        ]

        for idx, (label, val) in enumerate(properties):
            row = table.rows[idx]
            
            # 属性名单元格
            cell_lbl = row.cells[0]
            cell_lbl.width = Inches(1.8)
            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_lbl.paragraph_format.space_after = Pt(12)
            run_lbl = p_lbl.add_run(label)
            run_lbl.font.name = 'SimSun'
            run_lbl.font.hansi_name = 'SimSun'
            run_lbl.font.size = Pt(12) # 小四
            run_lbl.bold = True

            # 属性值单元格
            cell_val = row.cells[1]
            cell_val.width = Inches(3.7)
            p_val = cell_val.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_val.paragraph_format.space_after = Pt(12)
            
            # 使用下划线
            run_val = p_val.add_run(val if val else "                             ")
            run_val.font.name = 'SimSun'
            run_val.font.hansi_name = 'SimSun'
            run_val.font.size = Pt(12)
            run_val.underline = True

        # 添加分页符
        self.doc.add_page_break()

    def generate_task_sheet(self):
        """生成任务书页"""
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(10)
        p_title.paragraph_format.space_after = Pt(15)
        run_title = p_title.add_run("软件工程 专业软件开发综合实践课程设计任务书")
        run_title.font.name = 'SimHei'
        run_title.font.hansi_name = 'SimHei'
        run_title.font.size = Pt(15) # 三号
        run_title.bold = True

        # 9行6列的表格
        table = self.doc.add_table(rows=9, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table, "A0A0A0")

        # 列宽定义
        widths = [Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.1), Inches(0.8), Inches(1.2)]
        for row in table.rows:
            for c_idx, width in enumerate(widths):
                row.cells[c_idx].width = width

        # Row 0: 学生信息
        table.cell(0, 0).text = "学生姓名"
        table.cell(0, 1).text = ""
        table.cell(0, 2).text = "专业班级"
        table.cell(0, 3).text = "软件 230X 班"
        table.cell(0, 4).text = "学号"
        table.cell(0, 5).text = ""

        # 美化第一行字体
        for col_idx in range(6):
            cell = table.cell(0, col_idx)
            set_cell_background(cell, "F2F2F2") if col_idx % 2 == 0 else None
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'SimSun'
                r.font.hansi_name = 'SimSun'
                r.font.size = Pt(10)

        # Row 1: 题目
        table.cell(1, 0).text = "题    目"
        merge_cells_and_set_text(table, 1, 1, 1, 5, "基于Go语言与云原生微服务架构的社交平台（Twitter克隆版）的设计与实现", align_center=False, bold=True, size=10)
        set_cell_background(table.cell(1, 0), "F2F2F2")
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 2: 课题性质 / 课题来源
        table.cell(2, 0).text = "课题性质"
        table.cell(2, 1).text = "工程设计"
        table.cell(2, 2).text = "课题来源"
        merge_cells_and_set_text(table, 2, 3, 2, 5, "自拟课题", align_center=False, bold=False, size=10)
        for c in [0, 2]:
            set_cell_background(table.cell(2, c), "F2F2F2")
            table.cell(2, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 3: 指导教师 / 同组姓名
        table.cell(3, 0).text = "指导教师"
        table.cell(3, 1).text = ""
        table.cell(3, 2).text = "同组姓名"
        merge_cells_and_set_text(table, 3, 3, 3, 5, "", align_center=False, bold=False, size=10)
        for c in [0, 2]:
            set_cell_background(table.cell(3, c), "F2F2F2")
            table.cell(3, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Helper to populate long rows
        def fill_section_row(row_idx, label, text):
            set_cell_background(table.cell(row_idx, 0), "F2F2F2")
            table.cell(row_idx, 0).text = label
            table.cell(row_idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_lbl = table.cell(row_idx, 0).paragraphs[0]
            p_lbl.runs[0].font.size = Pt(10)
            p_lbl.runs[0].font.name = 'SimSun'
            p_lbl.runs[0].font.hansi_name = 'SimSun'
            
            cell_content = merge_cells_and_set_text(table, row_idx, 1, row_idx, 5, text, align_center=False, size=9)
            set_cell_margins(cell_content, top=120, bottom=120, left=150, right=150)

        # Row 4: 主要内容
        main_content_text = (
            "设计并实现一个基于微服务架构与云原生部署的仿 Twitter 社交系统。使用 Go 语言、gRPC、Consul、MySQL、Redis、RabbitMQ、MinIO、Kubernetes、ArgoCD、Temporal、ES与Qdrant等技术，主要完成以下开发与系统治理内容：\n"
            "1. 用户微服务：支持用户注册、登录、JWT 认证与 JWKS 非对称公钥验签、个人资料管理。\n"
            "2. 推文微服务：支持推文发布与多媒体上传、转发、书签、点赞、二级评论树，及大V混合 Feed 流（推拉结合与本地缓存防抖）。\n"
            "3. 关注微服务：支持用户关注与取关、粉丝列表、基于防抖门限的关注数据统计与大V变更维护。\n"
            "4. AI 智能体微服务：支持 RAG 双路召回语义搜索、可视化工作流 DSL 编排设计与 Temporal 状态机任务执行。\n"
            "5. 舆情与风控：基于分词与防刷模型的时间衰减热搜榜；支持影子风控与原子洗地防线。\n"
            "6. 前端系统：Vue 3 拖拽连线工作流编辑器与 Flutter 移动端 App 的双端设计与实现。\n"
            "7. 云原生交付与治理：通过 Helm 与 ArgoCD GitOps 交付部署；注入 Chaos Mesh 混沌故障，并利用 AIOps 大模型进行告警智能 RCA 诊断与规则自愈。"
        )
        fill_section_row(4, "主\n要\n内\n容", main_content_text)

        # Row 5: 任务要求
        req_text = (
            "1. 综合运用软件工程相关理论、方法与工具完成系统全生命周期的规范化实践。\n"
            "2. 完成系统可行性分析、核心业务流与领域模型分析，编写用例场景说明。\n"
            "3. 完成系统架构设计、数据表物理设计、主要模块职责划分与系统质量属性分析。\n"
            "4. 完成微服务高并发多级缓存、幂等事务、AI 编排引擎与双端适配的编码实现。\n"
            "5. 部署自动化 CI/CD 与 GitOps 交付链，搭建分布式追踪与可观测看板。\n"
            "6. 编写压力测试与混沌故障注入方案，并构建告警智能化自愈闭环。\n"
            "7. 提交格式规范、逻辑清晰的课程设计报告，能现场演示系统并对关键设计与决策进行答辩。"
        )
        fill_section_row(5, "任\n务\n要\n求", req_text)

        # Row 6: 参考文献
        ref_text = (
            "[1] 《软件工程：实践者的研究方法（原书第 9 版）》. [美] 罗杰 S. 普莱斯曼，布鲁斯 R. 马克西姆 著. 机械工业出版社. 2021.\n"
            "[2] 《软件架构实践（原书第 4 版）》. [美] 伦·巴斯，保罗·克莱门茨，瑞克·凯兹曼 著. 机械工业出版社. 2022.\n"
            "[3] 《Head First 设计模式（第二版）》. [美] 埃里克·弗里曼，伊丽莎白·罗布森 著. 中国电力出版社. 2022.\n"
            "[4] 《软件测试的艺术（原书第 3 版）》. [美] Glenford J. Myers，Corey Sandler，Tom Badgett 著. 机械工业出版社. 2023.\n"
            "[5] 《人月神话：软件项目管理之道（40 周年中文纪念版）》. [美] 小弗雷德里克·布鲁斯 著. 清华大学出版社. 2015."
        )
        fill_section_row(6, "参\n考\n文\n献", ref_text)

        # Row 7: 审查意见
        opinion_text = "同意\n\n\n\n教研室主任签字： 黄送       2026 年 5 月 26 日"
        fill_section_row(7, "审\n查\n意\n见", opinion_text)

        # Row 8: 说明栏
        merge_cells_and_set_text(table, 8, 0, 8, 5, "说明：本表由指导教师填写，由教研室主任审核后下达给选题学生，装订在设计（论文）首页.", align_center=False, size=9)
        table.rows[8].cells[0].paragraphs[0].runs[0].font.italic = True
        set_cell_margins(table.cell(8, 0), top=60, bottom=60, left=100, right=100)

        # 统一单元格居中及内边距样式
        for r_idx in range(4):
            for c_idx in range(6):
                cell = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

        # 添加分页符
        self.doc.add_page_break()

    def generate_toc_page(self):
        """生成目录占位页并包含Word自动目录域"""
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(20)
        p_title.paragraph_format.space_after = Pt(20)
        run_title = p_title.add_run("目 录")
        run_title.font.name = 'SimHei'
        run_title.font.hansi_name = 'SimHei'
        run_title.font.size = Pt(16) # 三号
        run_title.bold = True

        p_desc = self.doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc.paragraph_format.space_after = Pt(10)
        run_desc = p_desc.add_run("（请在 Word 中右键下方目录域并选择“更新域”生成目录）")
        run_desc.font.name = 'SimSun'
        run_desc.font.size = Pt(10.5)
        run_desc.font.italic = True
        run_desc.font.color.rgb = RGBColor(128, 128, 128)

        # 写入 Word 的 TOC 域代码
        p_toc = self.doc.add_paragraph()
        p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_toc.paragraph_format.space_after = Pt(20)
        
        run_toc = p_toc.add_run()
        # 域起始
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        # 域代码定义
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        # 域分割线
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        # 域结束
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_toc._r.append(fldChar1)
        run_toc._r.append(instrText)
        run_toc._r.append(fldChar2)
        run_toc._r.append(fldChar3)

        # 添加分页符
        self.doc.add_page_break()

    def parse_markdown_content(self):
        """解析 Markdown 并写入正文"""
        with open(self.md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []

        i = 0
        while i < len(lines):
            line = lines[i].rstrip('\n')
            stripped = line.strip()

            # 1. 忽略 Markdown 顶层标题，因为已经有了封面页
            if stripped.startswith('# ') and not stripped.startswith('##'):
                i += 1
                continue

            # 2. 代码块解析
            if stripped.startswith('```'):
                if in_code_block:
                    # 将缓存的代码块写入 Word
                    table = self.doc.add_table(rows=1, cols=1)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = False
                    table.columns[0].width = Inches(5.77)
                    cell = table.cell(0, 0)
                    set_cell_background(cell, 'F5F5F5')
                    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                    
                    # 隐藏左右边框，保留非常细的上下边框
                    tblPr = table._element.tblPr
                    tblBorders = OxmlElement('w:tblBorders')
                    for b_name in ['top', 'bottom']:
                        b = OxmlElement(f'w:{b_name}')
                        b.set(qn('w:val'), 'single')
                        b.set(qn('w:sz'), '4')
                        b.set(qn('w:color'), 'E0E0E0')
                        tblBorders.append(b)
                    for b_name in ['left', 'right']:
                        b = OxmlElement(f'w:{b_name}')
                        b.set(qn('w:val'), 'none')
                        tblBorders.append(b)
                    tblPr.append(tblBorders)

                    cp = cell.paragraphs[0]
                    cp.paragraph_format.line_spacing = 1.15
                    cp.paragraph_format.space_after = Pt(2)
                    
                    for idx, code_line in enumerate(code_content):
                        if idx > 0:
                            cp = cell.add_paragraph()
                            cp.paragraph_format.line_spacing = 1.15
                            cp.paragraph_format.space_after = Pt(2)
                        run = cp.add_run(code_line)
                        run.font.name = 'Consolas'
                        run.font.ascii_name = 'Consolas'
                        run.font.hansi_name = 'Consolas'
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                    
                    # 空行缓冲
                    p_space = self.doc.add_paragraph()
                    p_space.paragraph_format.space_after = Pt(4)
                    
                    in_code_block = False
                    code_content = []
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # 3. 表格解析
            if stripped.startswith('|'):
                in_table = True
                table_rows.append(stripped)
                i += 1
                continue
            elif in_table:
                in_table = False
                if len(table_rows) >= 2:
                    clean_rows = []
                    for tr in table_rows:
                        if re.match(r'^\|[\s:-|]*\|$', tr):
                            continue
                        clean_rows.append(tr)
                    
                    first_row_cells = [c.strip() for c in clean_rows[0].split('|')[1:-1]]
                    col_count = len(first_row_cells)
                    
                    table = self.doc.add_table(rows=len(clean_rows), cols=col_count)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(table, "D0D0D0")
                    
                    for row_idx, raw_row in enumerate(clean_rows):
                        cols = [c.strip() for c in raw_row.split('|')[1:-1]]
                        row = table.rows[row_idx]
                        for col_idx in range(min(col_count, len(cols))):
                            cell = row.cells[col_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.line_spacing = 1.15
                            
                            run = p.add_run(cols[col_idx])
                            run.font.name = 'SimSun'
                            run.font.ascii_name = 'Times New Roman'
                            run.font.hansi_name = 'SimSun'
                            
                            if row_idx == 0:
                                run.font.size = Pt(10)
                                run.bold = True
                                set_cell_background(cell, 'F2F2F2')
                            else:
                                run.font.size = Pt(9.5)
                            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                
                p_space = self.doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(4)
                table_rows = []

            # 4. 忽略空行
            if stripped == "":
                i += 1
                continue

            # 5. 一级标题转换（Markdown的 ## 转换为 Word的 Heading 1）
            h1_match = re.match(r'^##\s+(.*)', stripped)
            if h1_match:
                title = h1_match.group(1)
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.keep_with_next = True
                
                run = p.add_run(title)
                run.font.name = 'SimHei'
                run.font.hansi_name = 'SimHei'
                run.font.size = Pt(15) # 三号 (15-16pt)
                run.bold = True
                i += 1
                continue

            # 6. 二级标题转换（Markdown的 ### 转换为 Word的 Heading 2）
            h2_match = re.match(r'^###\s+(.*)', stripped)
            if h2_match:
                title = h2_match.group(1)
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                
                run = p.add_run(title)
                run.font.name = 'SimHei'
                run.font.hansi_name = 'SimHei'
                run.font.size = Pt(12) # 四号 (12pt)
                run.bold = True
                i += 1
                continue

            # 7. 无序列表转换（Markdown的 - 或 *）
            li_match = re.match(r'^[-*]\s+(.*)', stripped)
            if li_match:
                content = li_match.group(1)
                p = self.doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.25
                
                add_paragraph_with_runs(p, content)
                for r in p.runs:
                    if r.font.name != 'Consolas':
                        r.font.name = 'SimSun'
                        r.font.ascii_name = 'Times New Roman'
                        r.font.hansi_name = 'SimSun'
                    r.font.size = Pt(10.5)
                i += 1
                continue

            # 8. 有序列表转换（Markdown的 1. ）
            oli_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
            if oli_match:
                num = oli_match.group(1)
                content = oli_match.group(2)
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.25
                
                run_num = p.add_run(f"{num}. ")
                run_num.font.name = 'SimSun'
                run_num.font.ascii_name = 'Times New Roman'
                run_num.font.hansi_name = 'SimSun'
                run_num.font.size = Pt(10.5)
                
                add_paragraph_with_runs(p, content)
                for r in p.runs:
                    if r != run_num and r.font.name != 'Consolas':
                        r.font.name = 'SimSun'
                        r.font.ascii_name = 'Times New Roman'
                        r.font.hansi_name = 'SimSun'
                    r.font.size = Pt(10.5)
                i += 1
                continue

            # 9. 引用块转换（Markdown的 > ）
            quote_match = re.match(r'^＞?\s*>\s*(.*)', stripped) or re.match(r'^>\s*(.*)', stripped)
            if quote_match:
                content = quote_match.group(1)
                # 清除 github Alerts 标记
                content = re.sub(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*', '', content)
                
                table = self.doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.columns[0].width = Inches(5.77)
                cell = table.cell(0, 0)
                set_cell_background(cell, 'F9F9F9')
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                
                tblPr = table._element.tblPr
                tblBorders = OxmlElement('w:tblBorders')
                # 粗灰左边框提示，其它无边框
                l_border = OxmlElement('w:left')
                l_border.set(qn('w:val'), 'single')
                l_border.set(qn('w:sz'), '24') # 3pt
                l_border.set(qn('w:space'), '0')
                l_border.set(qn('w:color'), 'A0A0A0')
                tblBorders.append(l_border)
                for b_name in ['top', 'bottom', 'right']:
                    b = OxmlElement(f'w:{b_name}')
                    b.set(qn('w:val'), 'none')
                    tblBorders.append(b)
                tblPr.append(tblBorders)
                
                cp = cell.paragraphs[0]
                cp.paragraph_format.line_spacing = 1.25
                cp.paragraph_format.space_after = Pt(0)
                
                add_paragraph_with_runs(cp, content)
                for r in cp.runs:
                    if r.font.name != 'Consolas':
                        r.font.name = 'KaiTi'
                        r.font.ascii_name = 'Times New Roman'
                        r.font.hansi_name = 'KaiTi'
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = RGBColor(102, 102, 102)
                
                p_space = self.doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(4)
                
                i += 1
                continue

            # 10. 普通文本段落转换
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6) # 段后6磅
            p.paragraph_format.line_spacing = 1.25 # 1.25倍行距
            p.paragraph_format.first_line_indent = Inches(0.3) # 首行缩进约2字符
            
            add_paragraph_with_runs(p, stripped)
            for r in p.runs:
                if r.font.name != 'Consolas':
                    r.font.name = 'SimSun'
                    r.font.ascii_name = 'Times New Roman'
                    r.font.hansi_name = 'SimSun'
                r.font.size = Pt(10.5) # 五号
            i += 1

    def run(self):
        print(f"[*] Starting document generation...")
        self.generate_cover_page()
        print(f"[+] Cover page generated.")
        self.generate_task_sheet()
        print(f"[+] Task sheet generated.")
        self.generate_toc_page()
        print(f"[+] Table of contents page generated.")
        self.parse_markdown_content()
        print(f"[+] Markdown chapters parsed and written.")
        
        self.doc.save(self.docx_path)
        print(f"[Success] Success: Generated final Word report to: {self.docx_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base_dir, "docs", "COURSE_DESIGN_REPORT.md")
    docx_path = os.path.join(base_dir, "（6）课程设计报告（含任务书）_生成版.docx")
    
    generator = ReportGenerator(md_path, docx_path)
    generator.run()
