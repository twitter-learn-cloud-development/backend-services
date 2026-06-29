import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# 辅助函数：设置单元格背景色
# ==========================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

# ==========================================
# 辅助函数：设置单元格边框
# ==========================================
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # 细灰边框
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ==========================================
# 辅助函数：给段落添加内嵌代码高亮（简单粗体处理）
# ==========================================
def add_paragraph_with_runs(p, text, is_code_block=False):
    # 处理粗体 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_text = part[2:-2]
            run = p.add_run(clean_text)
            run.bold = True
        else:
            # 过滤行内代码 `code`
            subparts = re.split(r'(`.*?`)', part)
            for subpart in subparts:
                if subpart.startswith('`') and subpart.endswith('`'):
                    clean_sub = subpart[1:-1]
                    run = p.add_run(clean_sub)
                    run.font.name = 'Consolas'
                    run.font.ascii_name = 'Consolas'
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(199, 37, 78) # 红色高亮行内代码
                else:
                    run = p.add_run(subpart)

# ==========================================
# 核心转换器
# ==========================================
def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # 设置页面布局（A4，常规页边距）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0) # 2.54 cm
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25) # 3.18 cm
        section.right_margin = Inches(1.25)
        section.page_width = Inches(8.27) # A4 宽度
        section.page_height = Inches(11.69) # A4 高度

    # 读取 Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 状态标记
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    # 循环遍历行
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # ==========================================
        # 1. 代码块处理
        # ==========================================
        if stripped.startswith('```'):
            if in_code_block:
                # 结束代码块写入
                # 为保证美观，我们在 Word 中用一个单单元格表格来装载代码块，以带有灰底背景
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                # 设置宽度为 A4 文本区宽度
                table.columns[0].width = Inches(5.77)
                cell = table.cell(0, 0)
                set_cell_background(cell, 'F5F5F5') # 浅灰色背景
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                
                # 隐藏表格边框，只留细灰框或无框
                tblPr = table._element.tblPr
                tblBorders = OxmlElement('w:tblBorders')
                for b_name in ['top', 'left', 'bottom', 'right']:
                    b = OxmlElement(f'w:{b_name}')
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), '6') # ~0.75 pt
                    b.set(qn('w:space'), '0')
                    b.set(qn('w:color'), 'E0E0E0')
                    tblBorders.append(b)
                tblPr.append(tblBorders)

                # 将代码行填入该单元格
                cp = cell.paragraphs[0]
                cp.paragraph_format.line_spacing = 1.15
                cp.paragraph_format.space_after = Pt(2)
                
                for idx, code_line in enumerate(code_content):
                    if idx > 0:
                        cp = cell.add_paragraph()
                        cp.paragraph_format.line_spacing = 1.15
                        cp.paragraph_format.space_after = Pt(2)
                    
                    run = cp.add_run(code_line)
                    run.font.name = 'Consolas'
                    run.font.ascii_name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(51, 51, 51)
                
                # 空一行
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                
                in_code_block = False
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # ==========================================
        # 2. 表格处理
        # ==========================================
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # 表格结束，输出表格
            in_table = False
            if len(table_rows) >= 2: # 至少要有表头和分割线
                # 过滤掉 `|---|---|` 这种分割线
                clean_rows = []
                for tr in table_rows:
                    if re.match(r'^\|[\s:-|]*\|$', tr):
                        continue
                    clean_rows.append(tr)
                
                # 解析列数
                first_row_cells = [c.strip() for c in clean_rows[0].split('|')[1:-1]]
                col_count = len(first_row_cells)
                
                table = doc.add_table(rows=len(clean_rows), cols=col_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                for row_idx, raw_row in enumerate(clean_rows):
                    cols = [c.strip() for c in raw_row.split('|')[1:-1]]
                    row = table.rows[row_idx]
                    for col_idx in range(min(col_count, len(cols))):
                        cell = row.cells[col_idx]
                        cell.text = "" # 清空默认段落
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.line_spacing = 1.15
                        
                        run = p.add_run(cols[col_idx])
                        run.font.name = 'SimSun'
                        run.font.ascii_name = 'Times New Roman'
                        
                        if row_idx == 0:
                            # 表头格式
                            run.font.size = Pt(10.5) # 五号
                            run.bold = True
                            set_cell_background(cell, 'F2F2F2') # 灰色表头
                        else:
                            # 普通行格式
                            run.font.size = Pt(9.5) # 小五
                            
                        # 设置单元格微边距
                        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
            # 空一行
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            table_rows = []
            # 继续处理当前行，不加 i += 1，因为当前行不是表格行，需要正常解析

        # ==========================================
        # 3. 标题与段落处理
        # ==========================================
        if stripped == "":
            i += 1
            continue

        # 检查一级标题 #
        h1_match = re.match(r'^#\s+(.*)', stripped)
        if h1_match:
            title = h1_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title)
            run.font.name = 'SimHei' # 黑体
            run.font.ascii_name = 'Arial'
            run.font.size = Pt(16) # 三号
            run.bold = True
            i += 1
            continue

        # 检查二级标题 ##
        h2_match = re.match(r'^##\s+(.*)', stripped)
        if h2_match:
            title = h2_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title)
            run.font.name = 'SimHei'
            run.font.ascii_name = 'Arial'
            run.font.size = Pt(14) # 小三
            run.bold = True
            i += 1
            continue

        # 检查三级标题 ###
        h3_match = re.match(r'^###\s+(.*)', stripped)
        if h3_match:
            title = h3_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title)
            run.font.name = 'SimHei'
            run.font.ascii_name = 'Arial'
            run.font.size = Pt(12) # 四号
            run.bold = True
            i += 1
            continue

        # 检查无序列表 -
        li_match = re.match(r'^[-*]\s+(.*)', stripped)
        if li_match:
            content = li_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.25
            
            # 自定义字体
            add_paragraph_with_runs(p, content)
            for r in p.runs:
                if r.font.name != 'Consolas':
                    r.font.name = 'SimSun'
                    r.font.ascii_name = 'Times New Roman'
                r.font.size = Pt(10.5)
            i += 1
            continue

        # 检查有序列表 1.
        oli_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if oli_match:
            num = oli_match.group(1)
            content = oli_match.group(2)
            # 有序列表我们可以直接作为段落处理，并手动加上前缀，免去样式嵌套的烦恼
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.25
            
            run_num = p.add_run(f"{num}. ")
            run_num.font.name = 'SimSun'
            run_num.font.ascii_name = 'Times New Roman'
            run_num.font.size = Pt(10.5)
            
            add_paragraph_with_runs(p, content)
            for r in p.runs:
                if r != run_num and r.font.name != 'Consolas':
                    r.font.name = 'SimSun'
                    r.font.ascii_name = 'Times New Roman'
                r.font.size = Pt(10.5)
            i += 1
            continue

        # 检查引用块 >
        quote_match = re.match(r'^＞?\s*>\s*(.*)', stripped) or re.match(r'^>\s*(.*)', stripped)
        if quote_match:
            content = quote_match.group(1)
            # 过滤一些 GitHub 警示框前缀
            content = re.sub(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*', '', content)
            
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Inches(5.77)
            cell = table.cell(0, 0)
            set_cell_background(cell, 'F9F9F9')
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # 设置左侧粗灰边框做提示，其他框隐藏
            tblPr = table._element.tblPr
            tblBorders = OxmlElement('w:tblBorders')
            # 左侧框
            l_border = OxmlElement('w:left')
            l_border.set(qn('w:val'), 'single')
            l_border.set(qn('w:sz'), '24') # 3pt 粗线
            l_border.set(qn('w:space'), '0')
            l_border.set(qn('w:color'), 'A0A0A0')
            tblBorders.append(l_border)
            # 其他三边设置为空
            for b_name in ['top', 'bottom', 'right']:
                b = OxmlElement(f'w:{b_name}')
                b.set(qn('w:val'), 'none')
                tblBorders.append(b)
            tblPr.append(tblBorders)
            
            cp = cell.paragraphs[0]
            cp.paragraph_format.line_spacing = 1.25
            cp.paragraph_format.space_after = Pt(0)
            
            add_paragraph_with_runs(cp, content)
            for r in cp.runs:
                if r.font.name != 'Consolas':
                    r.font.name = 'KaiTi' # 引用使用楷体更美观
                    r.font.ascii_name = 'Times New Roman'
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(102, 102, 102)
            
            # 空一行
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6) # 段后 6 磅
        p.paragraph_format.line_spacing = 1.25 # 1.25 倍行距
        p.paragraph_format.first_line_indent = Inches(0.3) # 首行缩进约 2 字符
        
        add_paragraph_with_runs(p, stripped)
        for r in p.runs:
            if r.font.name != 'Consolas':
                r.font.name = 'SimSun' # 宋体
                r.font.ascii_name = 'Times New Roman'
            r.font.size = Pt(10.5) # 五号
        i += 1

    # 保存文档
    doc.save(docx_path)
    print(f"Success: Generated Word report to {docx_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base_dir, "docs", "COURSE_DESIGN_REPORT.md")
    docx_path = os.path.join(base_dir, "（6）课程设计报告（含任务书）_生成版.docx")
    
    print(f"Reading Markdown from: {md_path}")
    print(f"Target Docx Path: {docx_path}")
    
    if os.path.exists(md_path):
        markdown_to_docx(md_path, docx_path)
    else:
        print("Error: docs/COURSE_DESIGN_REPORT.md does not exist.")

